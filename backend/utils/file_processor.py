import pandas as pd
import numpy as np
import os
from docx import Document
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """文件处理器"""
    
    def __init__(self):
        self.supported_formats = ['.xlsx', '.xls', '.docx']
    
    def process_file(self, file_path):
        """处理上传的文件"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext in ['.xlsx', '.xls']:
                return self._process_excel(file_path)
            elif file_ext == '.docx':
                return self._process_word(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
                
        except Exception as e:
            logger.error(f"文件处理错误: {str(e)}")
            raise
    
    def _process_excel(self, file_path):
        """处理Excel文件"""
        try:
            # 读取Excel文件
            df = pd.read_excel(file_path)
            
            # 获取列名
            columns = df.columns.tolist()
            
            # 获取前10行数据预览
            preview_data = df.head(10).to_dict('records')
            
            return {
                'success': True,
                'columns': columns,
                'preview': preview_data,
                'total_rows': len(df),
                'file_type': 'excel'
            }
            
        except Exception as e:
            logger.error(f"Excel文件处理错误: {str(e)}")
            raise
    
    def _process_word(self, file_path):
        """处理Word文件"""
        try:
            doc = Document(file_path)
            
            # 提取所有段落文本
            texts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    texts.append(paragraph.text.strip())
            
            # 创建DataFrame
            df = pd.DataFrame({
                'text': texts,
                'paragraph_id': range(len(texts))
            })
            
            # 获取列名
            columns = df.columns.tolist()
            
            # 获取前10行数据预览
            preview_data = df.head(10).to_dict('records')
            
            return {
                'success': True,
                'columns': columns,
                'preview': preview_data,
                'total_rows': len(df),
                'file_type': 'word'
            }
            
        except Exception as e:
            logger.error(f"Word文件处理错误: {str(e)}")
            raise
    
    def extract_texts_from_data(self, file_path, text_column, timestamp_column=None, file_type='excel'):
        """从文件中提取完整文本数据"""
        try:
            # 重新读取完整文件数据
            if file_type == 'excel':
                df = pd.read_excel(file_path)
            elif file_type == 'word':
                doc = Document(file_path)
                texts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        texts.append(paragraph.text.strip())
                df = pd.DataFrame({
                    'text': texts,
                    'paragraph_id': range(len(texts))
                })
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
            
            # 确保文本列存在
            if text_column not in df.columns:
                raise ValueError(f"列 '{text_column}' 不存在于文件中")
            
            # 提取文本，去除空值
            texts = df[text_column].dropna().astype(str).tolist()
            
            result = {
                'texts': texts,
                'total_documents': len(texts),
                'total_rows': len(df)
            }
            
            # 提取时间戳（如果指定）
            if timestamp_column and timestamp_column in df.columns:
                timestamps = df[timestamp_column].dropna().tolist()
                result['timestamps'] = timestamps
            
            return result
            
        except Exception as e:
            logger.error(f"文本提取错误: {str(e)}")
            raise
